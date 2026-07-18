"""Professional Word document builder for processed lessons."""

import re
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Pt, RGBColor, Emu

from core.model import (
    BookSettings,
    Lesson,
    ConversionStats,
    PAGE_DIMENSIONS_CM,
)
from utils.logger import get_logger

logger = get_logger(__name__)


class DocxBuilder:
    """Professional Word document builder.

    Converts a list of lessons into a formatted, print-ready DOCX file
    complete with a cover page, table of contents, page numbering, header, and footer.
    """

    def __init__(
        self,
        settings: BookSettings,
        progress_callback: Optional[Callable[[int, int, str], None]] = None,
        stop_check: Optional[Callable[[], bool]] = None,
    ) -> None:
        """Initializes the builder.

        Args:
            settings: Book settings configuration.
            progress_callback: Callback function invoked with (current, total, filename).
            stop_check: Callback function returning True if the user requests a stop.
        """
        self.settings = settings
        self._progress_callback = progress_callback
        self._stop_check = stop_check
        self.doc = Document()
        self._is_rtl = settings.text_direction == "RTL"

    # ══════════════════════════════════════════════
    #  Public Interface
    # ══════════════════════════════════════════════

    def build(self, lessons: list[Lesson]) -> ConversionStats:
        """Builds and saves the full document.

        Args:
            lessons: List of lessons to be included.

        Returns:
            Statistics regarding the conversion process.
        """
        stats = ConversionStats(total_files=len(lessons))
        start_time = datetime.now()

        try:
            self._setup_document()
            self._build_preliminary_pages(lessons)
            self._build_content(lessons, stats)
            self._finalize_document()
        except Exception as e:
            logger.error("Error during document building: %s", e, exc_info=True)
            stats.failed_files = stats.total_files - stats.processed_files
            raise

        stats.elapsed_seconds = (datetime.now() - start_time).total_seconds()
        logger.info(
            "Document built successfully: %d lessons in %.1f seconds",
            stats.processed_files,
            stats.elapsed_seconds,
        )
        return stats

    def save(self, output_path: str) -> None:
        """Saves the document to the specified destination path.

        Args:
            output_path: Output file path (must end with .docx).
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        logger.info("Document saved to: %s", output_path)

    # ══════════════════════════════════════════════
    #  Document Configuration
    # ══════════════════════════════════════════════

    def _setup_document(self) -> None:
        """Sets up the baseline configurations and the initial document section."""
        s = self.settings
        section = self.doc.sections[0]

        # Paper layout setup
        dims = PAGE_DIMENSIONS_CM.get(s.page_size, PAGE_DIMENSIONS_CM["A4"])
        section.page_width = Cm(dims[0])
        section.page_height = Cm(dims[1])

        # Page boundaries (margins)
        section.top_margin = Cm(s.margin_top_cm)
        section.bottom_margin = Cm(s.margin_bottom_cm)
        section.left_margin = Cm(s.margin_left_cm)
        section.right_margin = Cm(s.margin_right_cm)

        # Base orientation flow setup
        if self._is_rtl:
            self._set_section_rtl(section)

        # Clear default header and footer layouts for the initial section
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.clear()
        for p in section.footer.paragraphs:
            p.clear()

    def _set_section_rtl(self, section) -> None:
        """Enables right-to-left layout configuration for the given section."""
        sect_pr = section._sectPr
        bidi_elem = parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>')
        existing = sect_pr.find(qn("w:bidi"))
        if existing is not None:
            sect_pr.remove(existing)
        sect_pr.append(bidi_elem)

    # ══════════════════════════════════════════════
    #  Preliminary Pages (Cover & Copyright)
    # ══════════════════════════════════════════════

    def _build_preliminary_pages(self, lessons: list[Lesson]) -> None:
        """Generates the cover and copyright layouts inside the initial section."""
        s = self.settings

        if s.add_cover:
            self._add_cover_page()

        if s.add_copyright:
            self._add_copyright_page()

        # Isolate body text contents into a distinct structural section
        self._add_content_section()

    def _add_cover_page(self) -> None:
        """Appends the standard cover page layout."""
        s = self.settings

        # Vertical spacing to shift contents downwards toward the center
        for _ in range(6):
            self.doc.add_paragraph()

        # Decorative break line
        separator = self.doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run("─" * 40)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Document main title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(12)
        self._apply_rtl_to_paragraph(title_para)
        run = title_para.add_run(s.book_title or "عنوان الكتاب")
        run.font.size = Pt(36)
        run.font.name = s.heading_font_name
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        self._apply_rtl_to_run(run)

        # Decorative break line
        separator2 = self.doc.add_paragraph()
        separator2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = separator2.add_run("─" * 40)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Spacer
        self.doc.add_paragraph()

        # Author name layout
        author_para = self.doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_rtl_to_paragraph(author_para)
        run_a = author_para.add_run(s.author_name or "اسم المؤلف")
        run_a.font.size = Pt(22)
        run_a.font.name = s.font_name
        run_a.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        self._apply_rtl_to_run(run_a)

        # Calendar year indicator
        year_para = self.doc.add_paragraph()
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_y = year_para.add_run(str(datetime.now().year))
        run_y.font.size = Pt(16)
        run_y.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Page split marker
        self.doc.add_page_break()

    def _add_copyright_page(self) -> None:
        """Appends the legal copyright declaration layout."""
        s = self.settings
        copyright_text = s.copyright_text.format(
            year=datetime.now().year,
            author=s.author_name or "المؤلف",
            title=s.book_title or "الكتاب",
        )

        for _ in range(8):
            self.doc.add_paragraph()

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_rtl_to_paragraph(para)
        run = para.add_run(copyright_text)
        run.font.size = Pt(12)
        run.font.name = s.font_name
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        self._apply_rtl_to_run(run)

        self.doc.add_page_break()

    # ══════════════════════════════════════════════
    #  Main Content Section Setup
    # ══════════════════════════════════════════════

    def _add_content_section(self) -> None:
        """Appends a fresh isolated section for body elements containing distinct headers and running details."""
        s = self.settings

        new_section = self.doc.add_section()
        dims = PAGE_DIMENSIONS_CM.get(s.page_size, PAGE_DIMENSIONS_CM["A4"])
        new_section.page_width = Cm(dims[0])
        new_section.page_height = Cm(dims[1])
        new_section.top_margin = Cm(s.margin_top_cm)
        new_section.bottom_margin = Cm(s.margin_bottom_cm)
        new_section.left_margin = Cm(s.margin_left_cm)
        new_section.right_margin = Cm(s.margin_right_cm)

        if self._is_rtl:
            self._set_section_rtl(new_section)

        # Re-index progression numbers starting explicitly from page 1
        new_section.start_type = 2  # WD_SECTION.NEW_PAGE

        # Configure automatic page numbers
        if s.add_page_numbers:
            self._setup_page_numbers(new_section)

        # Configure dynamic heading running strings
        if s.add_header and s.header_text:
            self._setup_header(new_section, s.header_text)

        # Configure persistent running footers
        if s.add_footer and s.footer_text:
            self._setup_footer(new_section, s.footer_text)

    def _setup_page_numbers(self, section) -> None:
        """Injects automatic native page numbering tracking components into the footer field."""
        footer = section.footer
        footer.is_linked_to_previous = False

        # Verify a baseline workspace paragraph structure exists
        if not footer.paragraphs:
            footer.add_paragraph()
        para = footer.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Assemble structural XML fields mapping directly to standard page number fields
        run1 = para.add_run()
        fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1._r.append(fld_begin)

        run2 = para.add_run()
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._r.append(instr)

        run3 = para.add_run()
        fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_end)

    def _setup_header(self, section, text: str) -> None:
        """Configures the top header text parameters."""
        header = section.header
        header.is_linked_to_previous = False
        if not header.paragraphs:
            header.add_paragraph()
        para = header.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_rtl_to_paragraph(para)

        run = para.add_run(text)
        run.font.size = Pt(9)
        run.font.name = self.settings.font_name
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        self._apply_rtl_to_run(run)

        # Append a structural boundary accent directly beneath the running text
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        existing_bdr = pPr.find(qn("w:pBdr"))
        if existing_bdr is not None:
            pPr.remove(existing_bdr)
        pPr.append(pBdr)

    def _setup_footer(self, section, text: str) -> None:
        """Configures the bottom running footer strings alongside automatic page tags."""
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # If a field tracking element sequence exists, reassemble cleanly
        if para.runs:
            pass

        # Clear and reconstruct footer layout cleanly: textual elements + sequential digits
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_rtl_to_paragraph(para)

        if text:
            tr = para.add_run(text + "  ")
            tr.font.size = Pt(9)
            tr.font.name = self.settings.font_name
            tr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            self._apply_rtl_to_run(tr)

        # Setup sequence indexing elements via system XML tokens
        r1 = para.add_run()
        r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2 = para.add_run()
        r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r3 = para.add_run()
        r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    # ══════════════════════════════════════════════
    #  Content Sequencing (TOC, Intro, Main, Outro)
    # ══════════════════════════════════════════════

    def _build_content(self, lessons: list[Lesson], stats: ConversionStats) -> None:
        """Sequences core textual structural properties cleanly."""
        s = self.settings

        # Table of contents index page layout
        if s.add_toc and lessons:
            self._add_table_of_contents()
            self.doc.add_page_break()

        # Introduction block layout
        if s.intro_text.strip():
            self._add_intro_outro(s.intro_text, is_intro=True)

        # Core file content rendering loop
        for i, lesson in enumerate(lessons):
            # Intercept iteration if cancellation state is active
            if self._stop_check and self._stop_check():
                stats.cancelled = True
                logger.info("Operational execution terminated by caller statement.")
                return

            try:
                self._add_lesson(lesson, index=i)
                stats.processed_files += 1
            except Exception as e:
                stats.failed_files += 1
                stats.failed_filenames.append(lesson.title)
                logger.error("Failed to append structural segment elements '%s': %s", lesson.title, e)

            # Fire tracking state updates outward
            if self._progress_callback:
                self._progress_callback(i + 1, len(lessons), lesson.title)

        # Outro block layout
        if s.conclusion_text.strip():
            self._add_intro_outro(s.conclusion_text, is_intro=False)

    def _add_table_of_contents(self) -> None:
        """Appends a word-native structural Table of Contents tracking component field."""
        s = self.settings

        # Index page section label layout
        heading = self.doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_rtl_to_paragraph(heading)
        run = heading.add_run("جدول المحتويات")
        run.font.name = s.heading_font_name
        run.font.size = Pt(s.heading_font_size)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        self._apply_rtl_to_run(run)

        # XML field elements mapped to standard index tracking properties
        para = self.doc.add_paragraph()
        self._apply_rtl_to_paragraph(para)

        r1 = para.add_run()
        r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))

        r2 = para.add_run()
        r2._r.append(
            parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve">'
                f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
            )
        )

        r3 = para.add_run()
        r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))

        # Standard system text indicator telling readers how to render missing index details manually
        r4 = para.add_run("[ يُرجى فتح المستند في Word والضغط على Ctrl+A ثم F9 لتحديث الفهرس ]")
        r4.font.size = Pt(10)
        r4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r4.font.italic = True
        self._apply_rtl_to_run(r4)

        r5 = para.add_run()
        r5._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    def _add_intro_outro(self, text: str, is_intro: bool) -> None:
        """Appends basic isolated content sections for intros or outros."""
        s = self.settings
        title = "المقدمة" if is_intro else "الخاتمة"

        heading = self.doc.add_heading(level=1)
        self._apply_rtl_to_paragraph(heading)
        run = heading.add_run(title)
        run.font.name = s.heading_font_name
        run.font.size = Pt(s.heading_font_size)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        self._apply_rtl_to_run(run)

        # Parse character block lines into formatted layout segments
        paragraphs = re.split(r"\n\s*\n", text.strip())
        for para_text in paragraphs:
            unified = " ".join(para_text.split())
            if unified.strip():
                self._add_formatted_paragraph(unified.strip())

        self.doc.add_page_break()

    def _add_lesson(self, lesson: Lesson, index: int) -> None:
        """Appends an individual distinct structural text segment component block."""
        s = self.settings

        # Render heading text segment label (Heading 1 standard)
        heading = self.doc.add_heading(level=1)
        self._apply_rtl_to_paragraph(heading)
        run = heading.add_run(lesson.title)
        run.font.name = s.heading_font_name
        run.font.size = Pt(s.heading_font_size)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        self._apply_rtl_to_run(run)

        # Iterate individual text layout segments
        for para_text in lesson.paragraphs:
            self._add_formatted_paragraph(para_text)

        # Add section breaks to cleanly split document bodies
        self.doc.add_page_break()

    def _add_formatted_paragraph(self, text: str) -> None:
        """Creates an aligned text block conforming strictly to styling attributes.

        Args:
            text: Pure content character string to be appended.
        """
        s = self.settings
        para = self.doc.add_paragraph()
        self._apply_rtl_to_paragraph(para)

        # Alignment adjustments based on bi-directional setup
        if self._is_rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Bottom trailing paragraph margins
        para.paragraph_format.space_after = Cm(s.paragraph_spacing_after_cm)

        # Height lines spacing scaling metrics
        para.paragraph_format.line_spacing = s.line_spacing

        # Indentation offset rules for initial sentence strings
        if s.first_line_indent_cm > 0:
            para.paragraph_format.first_line_indent = Cm(s.first_line_indent_cm)

        # Commit character strings safely
        run = para.add_run(text)
        run.font.size = Pt(s.font_size)
        run.font.name = s.font_name
        self._apply_rtl_to_run(run)

    # ══════════════════════════════════════════════
    #  Cleanups & Normalization Procedures
    # ══════════════════════════════════════════════

    def _finalize_document(self) -> None:
        """Performs optimization and styling updates before finalizing execution files."""
        # Strip structural trailing blank breaks
        body = self.doc.element.body
        paragraphs = body.findall(qn("w:p"))
        if paragraphs:
            last_para = paragraphs[-1]
            for run_elem in last_para.findall(qn("w:r")):
                for br in run_elem.findall(qn("w:br")):
                    br_type = br.get(qn("w:type"))
                    if br_type == "page" or br_type is None:
                        run_elem.remove(br)

        # Apply fallback system baseline default styles
        self._update_default_styles()

    def _update_default_styles(self) -> None:
        """Updates standard underlying typography attributes based on runtime user targets."""
        s = self.settings

        # Configure Normal paragraph baseline profiles
        try:
            normal_style = self.doc.styles["Normal"]
            normal_style.font.name = s.font_name
            normal_style.font.size = Pt(s.font_size)
            normal_style.paragraph_format.line_spacing = s.line_spacing

            # Inject target properties for right-to-left layout maps
            rPr = normal_style.element.get_or_add_rPr()
            rtl_elem = rPr.find(qn("w:rtl"))
            if rtl_elem is None:
                rtl_elem = parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>')
                rPr.append(rtl_elem)

            # Assign complex character script markers for Arabic text layout engine maps
            cs_elem = rPr.find(qn("w:cs"))
            if cs_elem is None:
                cs_elem = parse_xml(
                    f'<w:cs {nsdecls("w")} w:val="1"/>'
                )
                rPr.append(cs_elem)

            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:cs"), s.font_name)
            rFonts.set(qn("w:ascii"), s.font_name)
            rFonts.set(qn("w:hAnsi"), s.font_name)
        except Exception as e:
            logger.warning("Could not adjust structural baseline profiles: %s", e)

        # Configure Heading 1 typography properties
        try:
            h1_style = self.doc.styles["Heading 1"]
            h1_style.font.name = s.heading_font_name
            h1_style.font.size = Pt(s.heading_font_size)
            h1_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
            h1_style.font.bold = True

            h1_rPr = h1_style.element.get_or_add_rPr()
            h1_rtl = h1_rPr.find(qn("w:rtl"))
            if h1_rtl is None:
                h1_rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>'))

            h1_rFonts = h1_rPr.find(qn("w:rFonts"))
            if h1_rFonts is None:
                h1_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                h1_rPr.insert(0, h1_rFonts)
            h1_rFonts.set(qn("w:cs"), s.heading_font_name)
            h1_rFonts.set(qn("w:ascii"), s.heading_font_name)
            h1_rFonts.set(qn("w:hAnsi"), s.heading_font_name)
        except Exception as e:
            logger.warning("Could not optimize custom section header layouts: %s", e)

    # ══════════════════════════════════════════════
    #  Bi-directional Language Alignment Handlers
    # ══════════════════════════════════════════════

    def _apply_rtl_to_paragraph(self, paragraph) -> None:
        """Enforces a right-to-left directional state tracking matrix on a paragraph element."""
        if not self._is_rtl:
            return
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="1"/>'))

    def _apply_rtl_to_run(self, run) -> None:
        """Enforces a right-to-left directional state tracking matrix on a run text sequence."""
        if not self._is_rtl:
            return
        rPr = run._r.get_or_add_rPr()
        rtl = rPr.find(qn("w:rtl"))
        if rtl is None:
            rPr.append(parse_xml(f'<w:rtl {nsdecls("w")} w:val="1"/>'))